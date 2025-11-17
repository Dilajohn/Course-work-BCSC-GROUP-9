#QUESTION (4-5)=== COURSEWORK 2
"""  
AuthorS: OKUJA EMMANUEL DILA JOHN   2500728777  25/U//28777/PSA
         WASSWA KATEREGGA MAURICE   2500703613  25/U/03613/PSA
         NANFUKA JUSTINE            2500703528  25/U/03528/PS
         WAMIMBI CHRISTIAN          2500730993  25/U/30993/PSA

""" 
# Date: November 17, 2025

import math
import json
from datetime import datetime


ALLOWED_STUDENTS = {
    "OKUJA EMMANUEL DILA JOHN": "2500728777",
    "WASSWA KATEREGGA MAURICE": "2500703613"
}

class Calculator:
    def __init__(self):
        self.history = []                    # iii) List for history
        self.load_history()

        # iii) Dictionary maps symbols to functions → easy to extend!
        self.operations = {
            '+': self.add,
            '-': self.subtract,
            '*': self.multiply,
            '/': self.divide,
            '**': self.power,        # Exponentiation
            'sqrt': self.sqrt,       # Square root
            'sin': self.sin,
            'cos': self.cos,
            'tan': self.tan
        }

    # iv) Separate function for each operation
    def add(self, a, b=0):      return a + b
    def subtract(self, a, b=0): return a - b
    def multiply(self, a, b=0): return a * b
    def divide(self, a, b=0):
        if b == 0:
            print("Error: Division by zero!")
            return None
        return a / b
    def power(self, a, b=0):     return a ** b
    def sqrt(self, a, b=0):
        if a < 0:
            print("Error: Cannot take sqrt of negative!")
            return None
        return math.sqrt(a)
    def sin(self, a, b=0):       return math.sin(math.radians(a))
    def cos(self, a, b=0):       return math.cos(math.radians(a))
    def tan(self, a, b=0):       return math.tan(math.radians(a))

    # File management - saves history permanently
    def save_history(self):
        try:
            with open("calculator_history.json", "w") as f:
                json.dump(self.history, f)
        except:
            pass

    def load_history(self):
        try:
            with open("calculator_history.json", "r") as f:
                self.history = json.load(f)
        except:
            self.history = []

    def add_to_history(self, text):
        time = datetime.now().strftime("%d-%m-%Y %I:%M %p")
        entry = f"{time} → {text}"
        self.history.append(entry)
        self.save_history()

    def show_history(self):
        if not self.history:
            print("\nNo history yet.\n")
        else:
            print("\nYOUR CALCULATION HISTORY")
            print("-" * 50)
            for h in self.history[-10:]:
                print(h)
            print("-" * 50 + "\n")

    # Main program with login and 3 options only
    def start(self):
        print("ADVANCED CALCULATOR - UGANDA STUDENT VERSION")
        print("=" * 50)

        # Login
        name = input("\nEnter full name: ").strip().upper()
        stud_no = input("Enter student number: ").strip()

        if name in ALLOWED_STUDENTS and ALLOWED_STUDENTS[name] == stud_no:
            print(f"\nWelcome {name.title()}! Access Granted\n")
        else:
            print("\nIncorrect details. Goodbye!")
            return

        while True:
            print("MAIN MENU")
            print("1. Perform Calculation")
            print("2. View History")
            print("3. Exit")
            choice = input("\nChoose (1/2/3): ").strip()

            # Option 1: Perform Calculation
            if choice == "1":
                print("\nAvailable operations:")
                print("+  -  *  /  ** (power)  sqrt  sin  cos  tan")
                op = input("\nEnter operation: ").strip()

                # ii) Conditional check for valid operation
                if op not in self.operations:
                    print("Invalid operation!\n")
                    continue

                try:
                    if op in ['sqrt', 'sin', 'cos', 'tan']:
                        num = float(input("Enter number (degrees for trig): "))
                        result = self.operations[op](num)
                        expr = f"{op}({num})"
                    else:  # two-number operations
                        a = float(input("Enter first number: "))
                        b = float(input("Enter second number: "))
                        result = self.operations[op](a, b)
                        expr = f"{a} {op} {b}"

                    if result is not None:
                        print(f"\nResult: {result}\n")
                        self.add_to_history(f"{expr} = {result}")

                except ValueError:
                    print("Please enter valid numbers!\n")

            # Option 2: View History
            elif choice == "2":
                self.show_history()

            # Option 3: Exit
            elif choice == "3":
                print("Thank you for using the calculator!")
                break
            else:
                print("Please choose 1, 2 or 3 only!\n")


# Run the program
if __name__ == "__main__":
    calc = Calculator()
    calc.start()

"""
viii) You developed a CGPA software solution in your coursework. You are then tasked with
creating and adding a secure system form that captures the student details in that
application to authenticate users.
"""
# cgpa_secure_login.py
# Simple secure login system for your CGPA Calculator
import json
import hashlib
import getpass

# File where student accounts will be saved
USER_FILE = "cgpa_students.json"

# Simple password hashing (secure enough for school project)
def hash_password(password):
    return hashlib.sha256(password.encode()).hexdigest()

# Load existing students from file
def load_students():
    try:
        with open(USER_FILE, "r") as file:
            return json.load(file)
    except:
        return {}  # If file doesn't exist, start empty

# Save students to file
def save_students(students):
    with open(USER_FILE, "w") as file:
        json.dump(students, file, indent=4)

# Main Authentication Class 
class CGPAAuth:
    def __init__(self):
        self.students = load_students()

    # Register new student
    def register(self):
        print("\n=== REGISTER NEW STUDENT ===")
        reg_no = input("Enter Registration Number (e.g. 25/U/28777/PSA): ").strip()
        name = input("Enter Full Name: ").strip()

        if reg_no in self.students:
            print("This Reg. Number is already registered!")
            return False

        password = getpass.getpass("Create Password: ")
        confirm = getpass.getpass("Confirm Password: ")

        if password != confirm:
            print("Passwords do not match!")
            return False

        # Save securely with hashed password
        self.students[reg_no] = {
            "name": name,
            "password": hash_password(password)
        }
        save_students(self.students)
        print(f"Registration successful! Welcome {name}\n")
        return True

    # Login student
    def login(self):
        print("\n=== CGPA SYSTEM LOGIN ===")
        reg_no = input("Registration Number: ").strip()
        password = getpass.getpass("Password: ")

        if reg_no in self.students:
            if self.students[reg_no]["password"] == hash_password(password):
                print(f"\nLogin Successful! Welcome {self.students[reg_no]['name']}!")
                return reg_no  # Login success
            else:
                print("Wrong password!")
                return None
        else:
            print("Registration Number not found!")
            return None


if __name__ == "__main__":
    auth = CGPAAuth()

    while True:
        print("\n" + "="*40)
        print("   CGPA CALCULATOR - SECURE LOGIN")
        print("="*40)
        print("1. Register (First time users)")
        print("2. Login")
        print("3. Exit")
        choice = input("\nChoose option (1/2/3): ")

        if choice == "1":
            auth.register()
        elif choice == "2":
            user = auth.login()
            if user:
                print("\nYou now have access to the CGPA Calculator!")
                print("Your CGPA software is now unlocked")
                
                break
        elif choice == "3":
            print("Thank you. Goodbye!")
            break
        else:
            print("Please choose 1, 2 or 3 only!")

"""
ix) What is the output of the piece of program below once it's corrected, completed and
executed. Replace the student details with your own details
"""
class Student:
    def __init__(self, reg_number, name, stud_number):
        self.reg_number = reg_number
        self.name = name
        self.stud_number = stud_number 

    def print_student_details(self):
        print('name, regno, studeno:', self.name, self.reg_number, self.stud_number)

class CS_Student(Student):
    course = 'Computer Science'

    def print_student_details(self):
        super().print_student_details()  
        print('Course:', CS_Student.course)


student1 = Student('25/U/28777/PSA', 'OKUJA EMMANUEL DILA JOHN', 2500728777)
student2 = CS_Student('25/U/03613/PSA', 'WASSWA KATEREGGA MAURICE', 2500703613)

student1.print_student_details()
student2.print_student_details()


#QUESTION 5


from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

doc = Document()

# Title
title = doc.add_heading("COURSE WORK 1 AND 2", 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].font.size = Pt(18)
title.runs[0].bold = True

doc.add_paragraph("OKUJA EMMANUEL DILA JOHN\nStudent ID: 2500728777\nReg No: 25/U/28777/PSA\nDate: November 17, 2025", 
                  style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_page_break()

# QUESTION 1
doc.add_heading("QUESTION 1 – CGPA & BASIC CALCULATOR SYSTEM", level=1)

doc.add_heading("Semester 1 – CGPA: 5.0 (Distinction)", level=2)
t1 = doc.add_table(rows=5, cols=6)
t1.style = 'Table Grid'
hdr = t1.rows[0].cells
headers = ["Code", "Course Name", "Marks", "Grade", "GP", "CU"]
for i, h in enumerate(headers):
    hdr[i].text = h
    hdr[i].paragraphs[0].runs[0].bold = True

sem1 = [
    ["CSK 1101", "COMMUNICATION SKILLS", "89.0", "A", "5.0", "4"],
    ["CSC 1102", "STRUCTURED AND OBJECT ORIENTED PROGRAMMING", "86.0", "A", "5.0", "4"],
    ["CSC 1104", "COMPUTER ORGANIZATION AND ARCHITECTURE", "92.0", "A+", "5.0", "4"],
    ["CSC 1105", "MATHEMATICS", "97.0", "A+", "5.0", "4"]
]
for row in sem1:
    cells = t1.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

doc.add_heading("Semester 2 – CGPA: 5.0 (Distinction)", level=2)
t2 = doc.add_table(rows=5, cols=6)
t2.style = 'Table Grid'
for i, h in enumerate(headers):
    t2.rows[0].cells[i].text = h
    t2.rows[0].cells[i].paragraphs[0].runs[0].bold = True

sem2 = [
    ["CSC 1200", "OPERATING SYSTEMS", "82.0", "A", "5.0", "4"],
    ["CSC1201", "PROBABILITY AND STATISTICS", "93.0", "A+", "5.0", "4"],
    ["CSC 1202", "SOFTWARE DEVELOPMENT", "96.0", "A+", "5.0", "4"],
    ["CSC 1204", "DATA STRUCTURES AND ALGORITHMS", "87.0", "A", "5.0", "4"]
]
for row in sem2:
    cells = t2.add_row().cells
    for i, val in enumerate(row):
        cells[i].text = val

doc.add_paragraph("Formula Used: CGPA = Σ(GPᵢ × CUᵢ) / ΣCUᵢ\n", style='Intense Quote')

doc.add_heading("Basic Calculator Report", level=2)
doc.add_paragraph("Student Numbers: 216022204, 216002204, 216007570, 216002774")
doc.add_paragraph("Sum: 864034752")
doc.add_paragraph("Extracted CUs: CU1=64, CU2=03, CU3=47, CU4=52")
tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Table Grid'
tbl.cell(0,0).text = "Operation"; tbl.cell(0,1).text = "Result"
tbl.cell(0,0).paragraphs[0].runs[0].bold = True
tbl.cell(0,1).paragraphs[0].runs[0].bold = True
results = [("Addition", "166"), ("Subtraction", "-38"), ("Multiplication", "469248"), ("Division", "21.33")]
for i, (op, res) in enumerate(results, 1):
    tbl.cell(i,0).text = op
    tbl.cell(i,1).text = res

doc.add_page_break()

# QUESTION 4
doc.add_heading("QUESTION 4 – ADVANCED CALCULATOR WITH OOP", level=1)
parts = [
    
    "v)   Functions improve readability, reusability, testing, and maintainability.",
    "ix)  Output:\nname, regno, studeno: OKUJA EMMANUEL DILA JOHN 25/U/28777/PSA 2500728777\n"
         "name, regno, studeno: WASSWA KATEREGGA MAURICE 25/U/03613/PSA 2500703613\nCourse: Computer Science"
]

for part in parts:
    doc.add_paragraph(part)

doc.add_page_break()
doc.add_heading("GROUP MEMBERS", level=1)
members = [
    "OKUJA EMMANUEL DILA JOHN        2500728777     25/U/28777/PSA",
    "WASSWA KATEREGGA MAURICE         2500703613    25/U/03613/PSA",
    "NANFUKA JUSTINE                  2500703528    25/U/03528/PS",
    "WAMIMBI CHRISTIAN                2500730993    25/U/30993/PSA"
]
for m in members:
    doc.add_paragraph(m)

# Save
filename = "GROUP 9.docx"
doc.save(filename)
print(f"Python version generated: {os.path.abspath(filename)}")



"""
#c code
// question5_c_generator.c
#include <stdio.h>

int main() {
    FILE *f = fopen("GROUP 9.rtf", "w");
    fprintf(f, "{\\rtf1\\ansi\\deff0 {\\fonttbl {\\f0 Courier New;}}\\f0\\fs24\n");
    fprintf(f, "\\b\\fs32\\qc COURSE WORK 1 AND 2\\b0\\par\\par\n");
    fprintf(f, "\\qc OKUJA EMMANUEL DILA JOHN\\par 2500728777 | 25/U/28777/PSA\\par November 17, 2025\\par\\par\\par\n");

    fprintf(f, "\\b\\fs28 QUESTION 1 - CGPA RESULTS\\b0\\par\\par\n");
    fprintf(f, "Semester 1 & 2: CGPA 5.0 (Distinction)\\par\\par\n");
    fprintf(f, "Basic Calculator:\\par Sum: 864034752\\par CUs: 64, 03, 47, 52\\par ");
    fprintf(f, "Addition: 166 | Subtraction: -38 | Multiplication: 469248 | Division: 21.33\\par\\par\\par\n");

    fprintf(f, "\\b\\fs28 QUESTION 4 - ADVANCED CALCULATOR\\b0\\par\\par\n");
    fprintf(f, "i)   Modular structure with error handling\\par\n");
    fprintf(f, "ii)  if-else for validation\\par\n");
    fprintf(f, "iii) List + Dictionary mapping\\par\n");
    fprintf(f, "iv)  Separate functions for all operations\\par\n");
    fprintf(f, "v)   Functions improve reusability\\par\n");
    fprintf(f, "vi)  JSON file persistence\\par\n");
    fprintf(f, "vii) Full OOP Calculator class\\par\n");
    fprintf(f, "viii)Secure login with real student details\\par\n");
    fprintf(f, "ix)  Output: OKUJA... 2500728777 and WASSWA... Course: Computer Science\\par\\par\\par\n");

    fprintf(f, "\\b GROUP MEMBERS\\b0\\par\n");
    fprintf(f, "OKUJA EMMANUEL DILA JOHN        2500728777    25/U/28777/PSA\\par\n");
    fprintf(f, "WASSWA KATEREGGA MAURICE         2500703613    25/U/03613/PSA\\par\n");
    fprintf(f, "NANFUKA JUSTINE                  2500703528    25/U/03528/PS\\par\n");
    fprintf(f, "WAMIMBI CHRISTIAN                2500730993    25/U/30993/PSA\\par\\par\n");

    fprintf(f, "}\\par\n");
    fclose(f);
    printf("C version generated: FINAL_SUBMISSION.rtf\n");
    printf("Open with Microsoft Word - 100%% perfect!\n");
    return 0;
}

    """
