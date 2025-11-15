#QUESTION (1-3)=== COURSEWORK 1
""" 
AuthorS: OKUJA EMMANUEL DILA JOHN   2500728777  25/U//28777/PSA
         WASSWA KATEREGGA MAURICE   2500703613  25/U/03613/PSA
         NANFUKA JUSTINE            2500703528  25/U/03528/PS
         WAMIMBI CHRISTIAN          2500730993  25/U/30993/PSA 
""" 
# Date: November 15, 2025
# Purpose: CGPA Calculator + Basic Calculator with File I/O and Word Export

# QUESTION 1
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt
from typing import List, Dict, Tuple


GRADE_TO_GP = {
    'A+': 5.0, 'A': 5.0,
    'B+': 4.5,
    'B': 4.0,
    'C+': 3.5,
    'C': 3.0,
    'D+': 2.5,
    'D': 2.0,
    'E': 1.5,
    'E-': 1.0,
    'F': 0.0
}

MARKS_TO_GRADE = {
    (90, 100): 'A+',
    (80, 89): 'A',
    (75, 79): 'B+',
    (70, 74): 'B',
    (65, 69): 'C+',
    (60, 64): 'C',
    (55, 59): 'D+',
    (50, 54): 'D',
    (45, 49): 'E',
    (40, 44): 'E-',
    (0, 39): 'F'
}



class StudentCGPA:
    def __init__(self, student_id: str, name: str):
        self.student_id = student_id
        self.name = name
        self.history = []  # List of semester records
        self.load_history()

    

    def input_semester_data(self) -> Dict:
        print(f"\n--- Entering Data for {self.name} ({self.student_id}) ---")
        courses = []
        for i in range(4):
            print(f"\nCourse Unit {i+1}:")
            code = input("  Course Code: ").strip().upper()
            name = input("  Course Name: ").strip()
            marks = float(input("  Marks (0-100): ").strip())
            cu = int(input("  Credit Units: ").strip())
            grade = self.marks_to_grade(marks)
            gp = GRADE_TO_GP[grade]
            courses.append({
                'code': code,
                'name': name,
                'marks': marks,
                'grade': grade,
                'gp': gp,
                'cu': cu
            })
        return {'courses': courses, 'timestamp': datetime.now().strftime("%Y-%m-%d %H:%M")}

    

    def marks_to_grade(self, marks: float) -> str:
        for (low, high), grade in MARKS_TO_GRADE.items():
            if low <= marks <= high:
                return grade
        return 'F'

    def calculate_cgpa(self, semester_data: Dict) -> float:
        total_gp_cu = 0.0
        total_cu = 0
        for course in semester_data['courses']:
            total_gp_cu += course['gp'] * course['cu']
            total_cu += course['cu']
        return round(total_gp_cu / total_cu, 2) if total_cu > 0 else 0.0

    def classify_cgpa(self, cgpa: float) -> str:
        if cgpa >= 4.5:
            return "Distinction"
        elif cgpa >= 2.0:
            return "Pass"
        else:
            return "Fail"

    
    def get_filename(self) -> str:
        return f"cgpa_history_{self.student_id}.txt"

    def save_history(self):
        filename = self.get_filename()
        try:
            with open(filename, 'w') as f:
                for record in self.history:
                    f.write(f"{record['semester']}|{record['cgpa']}|{record['classification']}|{record['timestamp']}\n")
            print(f"CGPA history saved to {filename}")
        except Exception as e:
            print(f"Error saving file: {e}")

    def load_history(self):
        filename = self.get_filename()
        if os.path.exists(filename):
            try:
                with open(filename, 'r') as f:
                    for line in f:
                        parts = line.strip().split('|')
                        if len(parts) == 4:
                            self.history.append({
                                'semester': parts[0],
                                'cgpa': float(parts[1]),
                                'classification': parts[2],
                                'timestamp': parts[3]
                            })
                print(f"Loaded {len(self.history)} previous records.")
            except Exception as e:
                print(f"Error loading file: {e}. Starting fresh.")

    
    def add_semester(self, semester_name: str):
        data = self.input_semester_data()
        cgpa = self.calculate_cgpa(data)
        classification = self.classify_cgpa(cgpa)
        record = {
            'semester': semester_name,
            'cgpa': cgpa,
            'classification': classification,
            'timestamp': data['timestamp'],
            'courses': data['courses']
        }
        self.history.append(record)
        self.save_history()
        self.generate_word_report(record)
        print(f"\nCGPA for {semester_name}: {cgpa} → {classification}")

    
    def generate_word_report(self, record: Dict):
        doc = Document()
        doc.add_heading('CGPA Calculation Report', 0)
        doc.add_paragraph(f"Student ID: {self.student_id}")
        doc.add_paragraph(f"Name: {self.name}")
        doc.add_paragraph(f"Semester: {record['semester']}")
        doc.add_paragraph(f"Date: {record['timestamp']}")
        doc.add_paragraph(f"CGPA: {record['cgpa']} → {record['classification']}\n")

        # Table
        table = doc.add_table(rows=1, cols=6)
        hdr_cells = table.rows[0].cells
        headers = ['Code', 'Course Name', 'Marks', 'Grade', 'GP', 'CU']
        for i, h in enumerate(headers):
            hdr_cells[i].text = h
            hdr_cells[i].paragraphs[0].runs[0].font.bold = True

        for course in record['courses']:
            row = table.add_row().cells
            row[0].text = course['code']
            row[1].text = course['name']
            row[2].text = str(course['marks'])
            row[3].text = course['grade']
            row[4].text = str(course['gp'])
            row[5].text = str(course['cu'])

        # Formula
        doc.add_paragraph("\nFormula Used:")
        doc.add_paragraph("CGPA = Σ(GPᵢ × CUᵢ) / ΣCUᵢ", style='Intense Quote')

        filename = f"CGPA_Report_{self.student_id}_{record['semester'].replace(' ', '_')}.docx"
        try:
            doc.save(filename)
            print(f"Report saved as {filename}")
        except Exception as e:
            print(f"Error saving Word document: {e}")


def basic_calculator() -> Dict:
    # Student numbers
    students = [216022204, 216002204, 216007570, 216002774]
    total = sum(students)
    print(f"Sum of student numbers: {total}")

    # Extract CU values: remove leftmost '8' → 64034752 → CU1=64, CU2=03, CU3=47, CU4=52
    num_str = str(total).lstrip('8')  # Remove leading 8
    cu_values = [int(num_str[i:i+2]) for i in range(0, len(num_str), 2)]
    cu1, cu2, cu3, cu4 = cu_values[:4]

    operations = {
        'Addition': cu1 + cu2 + cu3 + cu4,
        'Subtraction': cu1 - cu2 - cu3 - cu4,
        'Multiplication': cu1 * cu2 * cu3 * cu4,
        'Division': round(cu1 / cu2, 2) if cu2 != 0 else "Error (div by 0)"
    }

    print("\nBasic Calculator Results:")
    for op, result in operations.items():
        print(f"  {op}: {result}")

    return {
        'student_numbers': students,
        'sum': total,
        'extracted_cus': [cu1, cu2, cu3, cu4],
        'operations': operations
    }


def main():
    print("=== CGPA & BASIC CALCULATOR SYSTEM ===\n")

    # --- Basic Calculator ---
    calc_results = basic_calculator()

    # --- CGPA Calculator ---
    print("\n" + "="*50)
    student_id = input("Enter your Student ID: ").strip()
    name = input("Enter your Name: ").strip()
    student = StudentCGPA(student_id, name)

    # Add two semesters
    for sem in ["Semester 1", "Semester 2"]:
        print(f"\n>>> Adding {sem} <<<")
        student.add_semester(sem)
        input("\nPress Enter to continue to next semester...")

    # Display history
    print("\n" + "="*50)
    print("CGPA HISTORY")
    for rec in student.history:
        print(f"{rec['semester']}: {rec['cgpa']} ({rec['classification']})")

    # Save basic calculator to Word too
    doc = Document()
    doc.add_heading('Basic Calculator Report', 0)
    doc.add_paragraph(f"Student Numbers: {', '.join(map(str, calc_results['student_numbers']))}")
    doc.add_paragraph(f"Sum: {calc_results['sum']}")
    doc.add_paragraph(f"Extracted CUs: CU1={calc_results['extracted_cus'][0]}, CU2={calc_results['extracted_cus'][1]}, "
                      f"CU3={calc_results['extracted_cus'][2]}, CU4={calc_results['extracted_cus'][3]}")
    doc.add_paragraph("\nResults:")
    for op, res in calc_results['operations'].items():
        doc.add_paragraph(f"• {op}: {res}")
    doc.save("Basic_Calculator_Report.docx")
    print("Basic calculator report saved as Basic_Calculator_Report.docx")

if __name__ == "__main__":
    main()


#QUESTION TWO 
"""
a) Initialize the structure members
struct course {
    char* cName;
    int cCode;
} one;

one.cName = "Computer Architecture";
one.cCode = 1104;

b) Allow user input via keyboard
#include <stdio.h>
#include <stdlib.h>

struct course {
    char* cName;
    int cCode;
} one;

int main() {
    char buffer[100];

    printf("Enter Course Name: ");
    fgets(buffer, sizeof(buffer), stdin);
    buffer[strcspn(buffer, "\n")] = 0;  
    one.cName = malloc(strlen(buffer) + 1);
    strcpy(one.cName, buffer);

    printf("Enter Course Code: ");
    scanf("%d", &one.cCode);

    return 0;
}

b) Allow user input via keyboard
#include <stdio.h>
#include <stdlib.h>

struct course {
    char* cName;
    int cCode;
} one;

int main() {
    char buffer[100];

    printf("Enter Course Name: ");
    fgets(buffer, sizeof(buffer), stdin);
    buffer[strcspn(buffer, "\n")] = 0;  // Remove newline
    one.cName = malloc(strlen(buffer) + 1);
    strcpy(one.cName, buffer);

    printf("Enter Course Code: ");
    scanf("%d", &one.cCode);

    return 0;
}

c) Declare and use pointer to access cCode

struct course* ptr = &one;
printf("Course Code via pointer: %d\n", ptr->cCode);

d) Function prototype for weThink (pass by value)

void weThink(struct course c);

"""



#QUESTION 3
#a) Handle division by zero in average calculation
def calculate_average(amount_collected: int, number_of_users: int) -> float:
    try:
        average = amount_collected / number_of_users
        return average
    except ZeroDivisionError:
        print("Error: Cannot divide by zero")
        return 0.0

#b) Handle missing file with try/except
try:
    with open("report.txt", "r") as file:
        content = file.read()
        print("File content:\n", content)
except FileNotFoundError:
    print("File not found.")

#c) OOP Concepts
"""i. Inheritance
Inheritance is a concept in Object-Oriented Programming (OOP) that allows a new class (the child or subclass) to
 inherit properties and behaviors from an existing class (the parent or superclass). This creates a hierarchical 
 "is-a" relationship, such as a "dog is an animal".

Importance in Large Applications

    *Code Reusability: It allows you to define common attributes and methods once in a parent class, and child 
    classes can automatically reuse them, which reduces code duplication and speeds up development.
    *Extensibility: New classes can be added easily by inheriting from existing ones, extending their functionality
     without modifying the original code.
    *Organization: It helps in structuring code in a logical, hierarchical manner, making large codebases easier 
    to understand and maintain. """

class Animal:
    # Parent class
    def __init__(self, name):
        self.name = name

    def speak(self):
        print("The animal makes a sound.")

class Dog(Animal):
    # Child class inheriting from Animal
    def speak(self):
        print(f"{self.name} barks.")

my_dog = Dog("Buddy")
my_dog.speak()  # Output: Buddy barks.

#the Dog class inherits from Animal, automatically getting the name attribute.
#  It then overrides the speak method to provide its own specific behavior. 

"""ii. Encapsulation 
Encapsulation is the practice of bundling an object's data (attributes) and the methods that operate on that data 
into a single unit, which is the class itself. A key aspect of encapsulation is data hiding, where the internal state
 of an object is protected from direct external access. Access to this data is managed through public methods, often
  called "getters" and "setters".

Importance in Large Applications

    *Data Integrity and Security: It protects an object's internal data from unauthorized or unintended changes,
      ensuring that the data remains in a valid state.
    *Modularity: Encapsulation hides the implementation details of a class, allowing developers to change the
     internal workings without affecting other parts of the program that use the class.
    *Simplified Maintenance: By creating a clear interface for a class, it becomes easier to maintain and debug
     code, as any changes to a class's internal logic will not affect the external code that uses it.""" 

class BankAccount:
    def __init__(self, balance):
        self.__balance = balance  # Private attribute using double underscore

    def deposit(self, amount):
        if amount > 0:
            self.__balance += amount
            print(f"Deposited: {amount}. New balance: {self.__balance}")

    def get_balance(self):
        return self.__balance

account = BankAccount(100)
# This will raise an error because __balance is private
# print(account.__balance)

account.deposit(50)  # Correct way to interact with the data
# Output: Deposited: 50. New balance: 150

#a double underscore (__) before an attribute name causes name mangling, making the attribute private and not
#  directly accessible from outside the class. The deposit method and get_balance method are the public interfaces
#  for interacting with the __balance data. 

"""iii. Polymorphism 
Polymorphism, which means "many forms," is the ability of an object to take on different forms or for a single
 action to be performed in different ways. It is often achieved through inheritance, where subclasses provide their
  own specific implementation of a method that is defined in their common superclass. 

Importance in Large Applications

    *Flexibility and Extensibility: It allows you to write generic code that can work with objects of different
     types, as long as they share a common interface. This makes it easy to add new classes without changing existing
      code.
    *Code Reusability: Polymorphism allows a single function or method to handle objects from multiple different
     classes, reducing the need for lengthy conditional statements (if/elif/else).
    *Simplified System Design: It enables you to program to an interface rather than a specific implementation,
     making your code more modular and loosely coupled. """


class Dog:
    def speak(self):
        return "Woof!"

class Cat:
    def speak(self):
        return "Meow!"

class Duck:
    def speak(self):
        return "Quack!"

def make_animal_speak(animal):
    # This function works for any object that has a 'speak' method
    print(animal.speak())

dog = Dog()
cat = Cat()
duck = Duck()

make_animal_speak(dog)   # Output: Woof!
make_animal_speak(cat)   # Output: Meow!
make_animal_speak(duck)  # Output: Quack!

#the make_animal_speak function can accept any object with a speak() method. The function doesn't need to know the
#  specific type of object passed to it; it simply calls the speak() method, and the appropriate implementation is 
# executed at runtime based on the object's type.  

#question 3 
#ii) Polymorphism in Python – Two Ways
# 1. Method Overriding (Runtime Polymorphism)
class Product:
    def __init__(self, name, price, stock):
        self.name = name
        self.price = price
        self.stock = stock

    def apply_discount(self):
        return self.price  # No discount

class Electronics(Product):
    def apply_discount(self):  # Override
        return self.price * 0.9   # 10% off

class Clothing(Product):
    def apply_discount(self):  # Override
        return self.price * 0.8   # 20% off

# 2. Duck Typing (No inheritance needed)
def apply_store_discount(item):
    return item.apply_discount()  # Works if object has method

# === Usage Example ===
products = [
    Electronics("Laptop", 500000, 5),
    Clothing("T-Shirt", 15000, 50),
    Product("Book", 8000, 100)
]

print("After Discounts:")
for p in products:
    print(f"{p.name}: {apply_store_discount(p):,.0f} UGX")
